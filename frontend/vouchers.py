import os
import tempfile
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLineEdit, QPushButton, QLabel,
    QMessageBox, QTextEdit
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QPixmap, QTextDocument, QTextCursor, QFont
from frontend.utils import http_get
from frontend.price.get_rates import convert_to_currency
from frontend.urls import API_URL_ORDERS
import re
import qrcode
from urllib.parse import quote
from docx.shared import Cm


class VouchersTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._setup_vouchers_tab()
        self._apply_styles()

    def _setup_vouchers_tab(self):
        main_layout = QVBoxLayout(self)
        main_layout.setAlignment(Qt.AlignTop)
        main_layout.setContentsMargins(25, 20, 25, 20)
        main_layout.setSpacing(25)

        #? ---------------- ENCABEZADO PRINCIPAL ----------------
        header_layout = QHBoxLayout()
        header_layout.setAlignment(Qt.AlignLeft)

        icon_label = QLabel()
        icon_label.setPixmap(
            QPixmap("frontend/icons/voucher.png").scaled(45, 45, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        )
        title_label = QLabel("Generar vales")
        title_label.setStyleSheet("""
            font-size: 32px;
            font-weight: 700;
            color: #222;
        """)

        header_layout.addWidget(icon_label)
        header_layout.addSpacing(10)
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        main_layout.addLayout(header_layout)

        line = QLabel()
        line.setFixedHeight(2)
        line.setStyleSheet("background-color: #ccc; margin-top: 5px; margin-bottom: 5px;")
        main_layout.addWidget(line)

        #? ---------------- ENTRADA Y BOT√ìN ----------------
        input_container = QVBoxLayout()
        input_container.setAlignment(Qt.AlignCenter)
        input_container.setSpacing(20)

        self.order_input = QLineEdit()
        self.order_input.setPlaceholderText("Ingrese los n√∫meros de orden (ej. 101, 102, 103)")
        self.order_input.setClearButtonEnabled(True)
        self.order_input.setFixedWidth(800)
        self.order_input.setStyleSheet("""
            QLineEdit {
                border: 2px solid #aaa;
                border-radius: 10px;
                padding: 14px;
                font-size: 16px;
                background-color: #fff;
            }
            QLineEdit:focus {
                border: 2px solid #333;
                background-color: #f9f9f9;
            }
        """)
        self.order_input.returnPressed.connect(self._generate_vouchers)

        self.generate_btn = QPushButton("Generar vales")
        self.generate_btn.setObjectName("primaryBtn")
        self.generate_btn.setFixedWidth(220)
        self.generate_btn.setCursor(Qt.PointingHandCursor)
        self.generate_btn.clicked.connect(self._generate_vouchers)

        input_container.addWidget(self.order_input, alignment=Qt.AlignCenter)
        input_container.addWidget(self.generate_btn, alignment=Qt.AlignCenter)

        main_layout.addLayout(input_container)
        main_layout.addStretch()

    def _generate_vouchers(self):
        order_text = self.order_input.text().strip()
        if not order_text:
            QMessageBox.warning(self, "Error", "Por favor ingrese al menos un n√∫mero de orden.")
            return

        processed_text = re.sub(r'[\n\s]+', ',', order_text)
        try:
            order_numbers = [int(num.strip()) for num in processed_text.split(',') if num.strip()]
        except ValueError:
            QMessageBox.warning(self, "Error", "Formato inv√°lido. Use solo n√∫meros separados por comas, espacios o nueva l√≠nea.")
            return

        if not order_numbers:
            QMessageBox.warning(self, "Error", "No se encontraron n√∫meros de orden v√°lidos.")
            return

        if len(order_numbers) > 3:
            QMessageBox.warning(self, "Advertencia", "M√°ximo 3 √≥rdenes por p√°gina. Se usar√°n las primeras 3.")
            order_numbers = order_numbers[:3]

        orders_data = []
        for order_id in order_numbers:
            order_data = self._get_order_data(order_id)
            if order_data:
                orders_data.append(order_data)
            else:
                QMessageBox.warning(self, "Error", f"No se pudo obtener la orden #{order_id}")

        if not orders_data:
            QMessageBox.warning(self, "Error", "No se pudieron obtener datos de ninguna orden.")
            return

        try:
            file_path = self._create_word_document(orders_data)
            QMessageBox.information(
                self, 
                "√âxito", 
                f"Documento actualizado exitosamente:\n{file_path}\n\n"
                f"√ìrdenes incluidas: {', '.join(str(o['idOrder']) for o in orders_data)}"
            )
            self.order_input.clear()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo generar el documento: {str(e)}")

    def _get_order_data(self, order_id):
        url = f"{API_URL_ORDERS}{order_id}/full_details/"
        r = http_get(url)
        if r and r.status_code == 200:
            return r.json()
        return None

    def _get_order_data(self, order_id):
        url = f"{API_URL_ORDERS}{order_id}/full_details/"
        r = http_get(url)
        if r and r.status_code == 200:
            return r.json()
        return None


    def _generate_qr_for_client(self, client_name: str, phone_reference: str):
        numero_whatsapp = "+5355352549"
        mensaje = (
            f"Hola MOE! Mi amigo {client_name} ({phone_reference}) me recomend√≥ su negocio! "
            f"Me gustar√≠a aprovechar el descuento del 10% en mi primera compra por su recomendaci√≥n ü§ç."
        )
        mensaje_codificado = quote(mensaje)
        enlace_whatsapp = f"https://wa.me/{numero_whatsapp}?text={mensaje_codificado}"

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=8,
            border=2,
        )
        qr.add_data(enlace_whatsapp)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")

        temp_dir = tempfile.gettempdir()
        file_name = f"qr_{client_name.replace(' ', '_')}_{phone_reference}.png"
        qr_path = os.path.join(temp_dir, file_name)
        img.save(qr_path)

        return qr_path

    def _create_voucher_image(self, order_data):
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setFixedSize(400, 800)
        text_edit.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 11px;
                margin: 0px;
            }
        """)
        
        client_name = order_data['client']['name']
        phone_ref = order_data['client']['phone_number']
        qr_path = self._generate_qr_for_client(client_name, phone_ref)
        if not os.path.exists(qr_path):
            qr_html = "<div style='color: red; font-weight: bold; font-size: 10px;'>‚ö†Ô∏è QR no encontrado</div>"
        else:
            absolute_qr_path = os.path.abspath(qr_path)
            qr_html = f"<img src='file:///{absolute_qr_path}' width='120' height='120' style='float: left; margin-right: 10px;' alt='QR Code'>"
        
        voucher_html = self._format_voucher_html(order_data, qr_html)
        text_edit.setHtml(voucher_html)
        text_edit.document().setDocumentMargin(8)
        pixmap = QPixmap(text_edit.size())
        text_edit.render(pixmap)
        temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        pixmap.save(temp_file.name, 'PNG', quality=100)
        temp_file.close()
        
        return temp_file.name

    def _format_voucher_html(self, order_data, qr_html):
        text = f"<div style='font-weight: bold; font-size: 12px;'>üî∞ ORDEN No. {order_data['idOrder']}</div><br>"
        text += f"üóì Fecha: {order_data['order_date']}<br>"
        text += f"üóì Entrega: {order_data['delivery_date']}<br>"

        for book_info in order_data['books']:
            book = book_info['book']
            additives = book_info['additives']
            discount = book_info['discount']
            cantidad = book_info['quantity']
            
            text += f"<div style='font-weight: bold;'>üìö T√≠tulo: {book.get('title', 'Desconocido')}</div><br>"
            text += f"üë§ Autor: {book.get('author', 'Sin autor')}<br>"
            
            if cantidad > 1:
                text += f"üî¢ Cantidad: {cantidad}<br>"

            caratula_name = "Tapa Normal"
            caratula_price = 0
            service_additives = []
            
            for additive in additives:
                if additive["name"].lower().startswith("car√°tula"):
                    caratula_name = additive['name']
                    caratula_price = additive['price']
                elif additive["name"].lower().startswith("servicio"):
                    service_additives.append(additive)

            base_price = book.get('number_pages', 0)
            precio_base_caratula = base_price + caratula_price
            
            cup_price_base = convert_to_currency(precio_base_caratula, 'USD', 'CUP')
            mlc_price_base = convert_to_currency(precio_base_caratula, 'USD', 'MLC')
            
            text += f"üí∞ {caratula_name}: {precio_base_caratula} USD<br>"

            for service in service_additives:
                service_cup = convert_to_currency(service['price'], 'USD', 'CUP')
                service_mlc = convert_to_currency(service['price'], 'USD', 'MLC')
                text += f"üí∞ {service['name']}: {service['price']} USD<br>"

            if discount != 0:
                text += f"üìâ Descuento: {discount}%<br>"

            precio_base_con_descuento = precio_base_caratula * (1 - discount / 100)
            precio_servicios = sum(service['price'] for service in service_additives)
            precio_unitario_final = precio_base_con_descuento + precio_servicios
            precio_total_libro = precio_unitario_final * cantidad
            
            if cantidad > 1:
                text += f"üí∞ Precio unitario: {precio_unitario_final:.2f} USD<br>"
                text += f"üí∞ Total libro: {precio_total_libro:.2f} USD<br><br>"
            else:
                text += f"üí∞ Total libro: {precio_total_libro:.2f} USD<br>"

        total_final = order_data['total_price']
        total_cup = convert_to_currency(total_final, 'USD', 'CUP')
        total_mlc = convert_to_currency(total_final, 'USD', 'MLC')
        text += f"<div style='font-weight: bold; font-size: 12px;'>üí∞ Total a pagar: {total_final:.2f} USD</div><br>"
        
        text += f"üí∞ Pago adelanto: {order_data['payment_advance']:.2f} USD<br>"
        text += f"üí∞ Pago pendiente: {order_data['outstanding_payment']:.2f} USD<br><br>"
        
        if order_data['delivery_price'] > 0:
            text += f"üöó Mensajer√≠a: {order_data['delivery_price']:.2f} USD<br>"
        else:
            text += f"üöó Mensajer√≠a: Recogida<br><br>"

        text += "<div style='font-weight: bold; font-size: 12px;'>üë§ Informaci√≥n de contacto:</div><br>"
        text += f"‚Äî Nombre: {order_data['client']['name']}<br>"
        text += f"‚Äî Carnet de identidad: {order_data['client']['identity']}<br>"
        text += f"‚Äî Tel√©fono: {order_data['client']['phone_number']}<br>"
        
        if (order_data.get('delivery_zone') and 
            'recogida' not in order_data['delivery_zone'].lower() and 
            order_data['address']):
            text += f"‚Äî Direcci√≥n de entrega: {order_data['address']}<br>"
        
        text += f"‚Äî Servicio de entrega: {order_data['_type']}<br>"
        text += f"‚Äî M√©todo de pago: {order_data['pay_method']}<br>"
         
        html_content = f"""
        <div style="font-family: 'Segoe UI', Arial, sans-serif; font-size: 11px; line-height: 1.4; margin: 0; padding: 0;">
            {text.replace('\n', '<br>')}
            <div style="margin-top: 10px; clear: both;">
                <div style="font-weight: bold; margin-bottom: 5px; font-size: 10px; text-align: left;">ü§ó √önete a nuestro WhatsApp</div>
                <div style="text-align: left;">
                    {qr_html}
                </div>
            </div>
        </div>
        """
        
        return html_content

    def _create_word_document(self, orders_data):
        documents_path  = os.path.join(os.path.expanduser("~"), "Documents")
        moe_path = os.path.join(documents_path, "Moe")
        os.makedirs(moe_path, exist_ok=True)


        filename = "vales_ordenes.docx"
        file_path = os.path.join(moe_path, filename)
        
        if os.path.exists(file_path):
            doc = Document(file_path)
        else:
            doc = Document()

        section = doc.sections[-1]
        section.orientation = 1
        section.page_width = Cm(27.94)
        section.page_height = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

        if len(doc.paragraphs) > 0:
            doc.add_page_break()

        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        col_width = Cm(9.33)
        for cell in table.rows[0].cells:
            cell.width = col_width

        temp_files = []
        try:
            for i, order_data in enumerate(orders_data):
                if i < 3:
                    image_path = self._create_voucher_image(order_data)
                    temp_files.append(image_path)
                    paragraph = table.rows[0].cells[i].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Cm(9.0))
        finally:
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except:
                    pass

        doc.save(file_path)
        return file_path

    def _apply_styles(self):
        self.setStyleSheet("""
            QPushButton#primaryBtn {
                background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #444, stop:1 #000);
                color: white;
                font-weight: 600;
                padding: 10px 20px;
                border-radius: 8px;
                border: 1px solid #111;
                font-size: 15px;
            }
            QPushButton#primaryBtn:hover {
                background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #555, stop:1 #111);
                border: 1px solid #222;
            }
            QPushButton#primaryBtn:pressed {
                background-color: #000;
                border: 1px solid #000;
                padding-top: 11px;
                padding-bottom: 9px;
            }
        """)